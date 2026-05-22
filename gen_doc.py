#!/usr/bin/env python3
# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

section = doc.sections[0]
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin   = Cm(3)
section.right_margin  = Cm(2)

def set_font(run, bold=False, size=12, color=None):
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = 'Times New Roman'
    if color:
        run.font.color.rgb = RGBColor(*color)

def heading1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    set_font(run, bold=True, size=16, color=(0x2C, 0x3E, 0x50))
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p

def heading2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    set_font(run, bold=True, size=14, color=(0x1A, 0x5E, 0x8A))
    return p

def heading3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    set_font(run, bold=True, size=12, color=(0x2E, 0x86, 0xAB))
    return p

def body(text, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.first_line_indent = Cm(1)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.italic = italic
    return p

def bullet(text):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    return p

def code_block(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F0F0F0')
    pPr.append(shd)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size  = Pt(9)
    return p

def diagram_placeholder(subject):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(10)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('<diagram for ' + subject + '>')
    run.font.name  = 'Arial'
    run.font.size  = Pt(11)
    run.italic     = True
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    return p

def add_table(headers, rows):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Table Grid'
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = 'Times New Roman'
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '2C3E50')
        tcPr.append(shd)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for ri, row_data in enumerate(rows):
        row = table.rows[ri+1]
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            cell.text = str(val)
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(9)
                run.font.name = 'Times New Roman'
    doc.add_paragraph()

def page_break():
    doc.add_page_break()

# ==============================================================================
# TITLE PAGE
# ==============================================================================

p = doc.add_paragraph()
p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(60)
run = p.add_run('ТЕХНИЧЕСКИ УНИВЕРСИТЕТ')
set_font(run, bold=True, size=16)

p = doc.add_paragraph()
p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Факултет по Компютърни системи и технологии')
set_font(run, bold=False, size=13)

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('ДИПЛОМНА РАБОТА')
set_font(run, bold=True, size=20, color=(0x1A, 0x5E, 0x8A))

doc.add_paragraph()

p = doc.add_paragraph()
p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('MyDrive')
set_font(run, bold=True, size=28, color=(0x2C, 0x3E, 0x50))

p = doc.add_paragraph()
p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Облачна платформа за съхранение и управление на документи')
set_font(run, bold=False, size=15, color=(0x2E, 0x86, 0xAB))

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Дипломант: Слави Овчаров Николаев')
set_font(run, bold=False, size=13)

p = doc.add_paragraph()
p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Специалност: Компютърни системи и технологии')
set_font(run, bold=False, size=13)

doc.add_paragraph()

p = doc.add_paragraph()
p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('2026 г.')
set_font(run, bold=False, size=13)

page_break()

# ==============================================================================
# TABLE OF CONTENTS
# ==============================================================================

heading1('СЪДЪРЖАНИЕ')
toc_items = [
    ('1.', 'ОПИСАНИЕ НА РЕШАВАНИЯ ПРОБЛЕМ'),
    ('  1.1.', 'Контекст и мотивация'),
    ('  1.2.', 'Съществуващи решения и техните ограничения'),
    ('  1.3.', 'Цели на системата'),
    ('  1.4.', 'Целева аудитория'),
    ('  1.5.', 'Обхват и ограничения'),
    ('2.', 'АРХИТЕКТУРА НА ПРОЕКТА И БАЗАТА ДАННИ'),
    ('  2.1.', 'Системна архитектура -- обзор'),
    ('  2.2.', 'Компоненти на системата'),
    ('  2.3.', 'Схема на базата данни'),
    ('  2.4.', 'Комуникация между компонентите'),
    ('  2.5.', 'Сигурност и автентикация'),
    ('  2.6.', 'Инфраструктура -- Docker и Nginx'),
    ('3.', 'ОПИСАНИЕ НА КОДА'),
    ('  3.1.', 'Backend -- Spring Boot'),
    ('  3.2.', 'File Server -- Go'),
    ('  3.3.', 'Frontend -- Next.js'),
    ('  3.4.', 'Реално-времево съвместно редактиране (WebSocket)'),
    ('  3.5.', 'Електронно подписване на документи (eIDAS PAdES)'),
    ('  3.6.', 'AI Чатбот'),
    ('4.', 'РЪКОВОДСТВО ЗА ПОТРЕБИТЕЛЯ'),
    ('  4.1.', 'Регистрация и вход в системата'),
    ('  4.2.', 'Управление на папки и файлове'),
    ('  4.3.', 'Качване и изтегляне на файлове'),
    ('  4.4.', 'Споделяне на файлове и папки'),
    ('  4.5.', 'Съвместно редактиране на документи'),
    ('  4.6.', 'Електронно подписване на PDF'),
    ('  4.7.', 'AI Асистент'),
    ('  4.8.', 'Управление на профила и абонамента'),
]
for num, title in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(num + '  ' + title)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    if not num.startswith('  '):
        run.bold = True

page_break()

# ==============================================================================
# CHAPTER 1
# ==============================================================================

heading1('1. ОПИСАНИЕ НА РЕШАВАНИЯ ПРОБЛЕМ')

heading2('1.1. Контекст и мотивация')

body(
    'В днешното дигитално общество организациите и отделните потребители генерират '
    'огромни обеми от документи -- договори, фактури, технически спецификации, '
    'изображения, архивни данни и много други. Традиционното съхранение на локален '
    'диск крие значителни рискове: загуба на данни вследствие на хардуерна повреда, '
    'липса на версиониране, затруднен споделен достъп и невъзможност за едновременна '
    'работа на повече от един потребител върху един и същ документ.'
)

body(
    'Облачните услуги за съхранение (Google Drive, Microsoft OneDrive, Dropbox) '
    'решават голяма част от тези проблеми, но представляват "черни кутии" -- '
    'потребителят не контролира нито инфраструктурата, нито спазването на регулаторни '
    'изисквания (напр. GDPR). За организации в регулирани сектори (финанси, право, '
    'здравеопазване) това е неприемливо.'
)

body(
    'MyDrive е разработена като отворена, самостоятелно управлявана (self-hosted) '
    'алтернатива, предоставяща всички ключови функции на съвременна облачна платформа -- '
    'съхранение на файлове, реално-времево съвместно редактиране, споделяне с управление '
    'на права и електронно подписване на документи по стандарта eIDAS PAdES.'
)

heading2('1.2. Съществуващи решения и техните ограничения')

add_table(
    ['Решение', 'Силни страни', 'Ограничения'],
    [
        ['Google Drive', 'Зрял продукт, богати функции', 'Данните се намират в Google; ограничено съответствие с GDPR'],
        ['Nextcloud', 'Self-hosted, отворен код', 'Сложна инсталация, ресурсоемко'],
        ['Dropbox', 'Лесен UX', 'Затворен код, скъп за организации'],
        ['SharePoint', 'Дълбока интеграция с MS 365', 'Висока цена, сложна администрация'],
        ['MyDrive', 'Self-hosted, eIDAS, реално-времево редактиране, AI асистент', 'В процес на разработка'],
    ]
)

heading2('1.3. Цели на системата')

body('Системата MyDrive цели постигане на следните основни задачи:')
bullet('Централизирано и сигурно съхранение на файлове и папки в облака.')
bullet('JWT-базирана автентикация с access и refresh токени, reset на парола по имейл.')
bullet('Споделяне на файлове и папки чрез генерирани токен-връзки с нива на достъп (VIEW / EDIT / DELETE) и дата на изтичане.')
bullet('Реално-времево съвместно редактиране на текстови документи чрез WebSocket (STOMP).')
bullet('Електронно подписване на PDF документи по стандарта PAdES-B-B, съвместим с eIDAS.')
bullet('AI-асистент (чатбот, интегриран с GPT-4o mini), достъпен от всяка страница.')
bullet('Абонаментни планове с различен капацитет на съхранение.')
bullet('Пълен контрол върху инфраструктурата чрез Docker контейнери и Nginx реверс прокси.')

heading2('1.4. Целева аудитория')

body('Системата е насочена към два вида потребители:')
bullet(
    'Крайни потребители -- служители, студенти, адвокати, счетоводители, '
    'всякакви специалисти, работещи ежедневно с документи.'
)
bullet(
    'Системни администратори -- отговарят за инсталацията, конфигурацията '
    'и поддръжката на MyDrive на собствена инфраструктура.'
)

heading2('1.5. Обхват и ограничения')

body(
    'MyDrive е уеб приложение, достъпно от всеки модерен браузър. '
    'Системата не включва собствен мобилен клиент (предвиден в бъдеща версия). '
    'Версията за дипломна защита използва demo самоподписан сертификат за eIDAS -- '
    'в производствена среда се изисква сертификат от квалифициран доставчик на '
    'удостоверителни услуги (QSP). '
    'Абонаментните планове включват само управление на лимит за съхранение; '
    'платежна интеграция не е реализирана в текущата версия.'
)

page_break()

# ==============================================================================
# CHAPTER 2
# ==============================================================================

heading1('2. АРХИТЕКТУРА НА ПРОЕКТА И БАЗАТА ДАННИ')

heading2('2.1. Системна архитектура -- обзор')

body(
    'MyDrive е изградена като многослойна уеб приложение (multi-tier web application), '
    'следваща принципа за разделение на отговорностите (Separation of Concerns). '
    'Системата се състои от четири логически компонента, описани в таблицата по-долу.'
)

add_table(
    ['Компонент', 'Технология', 'Роля', 'Порт'],
    [
        ['Java API (MYDrive)', 'Spring Boot 3.5, Java 21', 'Бизнес логика, автентикация, метаданни', '8083'],
        ['File Server', 'Go', 'Физическо съхранение и изтегляне на двоични файлове', '8080'],
        ['Web Frontend', 'Next.js 15, React 19, TypeScript', 'Потребителски интерфейс в браузъра', '3000'],
        ['MySQL', 'MySQL 8', 'Релационна база данни за метаданни', '3306'],
    ]
)

diagram_placeholder('system architecture overview')

body(
    'Nginx работи като реверс прокси пред всички компоненти. '
    'Той отговаря за SSL терминация, маршрутизиране на заявките към правилния бекенд '
    'и балансиране на натоварването при мащабиране.'
)

heading2('2.2. Компоненти на системата')

heading3('2.2.1. Spring Boot API (MYDrive)')

body(
    'Централният компонент на системата. Реализира REST API, обработва автентикация '
    'с JWT, управлява метаданните за файлове, папки, споделяния и потребители, '
    'оркестрира eIDAS подписването и комуникира с OpenAI за чатбот функционалността.'
)

body('Ключови пакети и класове:')
bullet('controllers/ -- HTTP контролери: AuthController, FileController, FolderController, ShareController, CollaborationController, DocumentSigningController.')
bullet('services/ -- бизнес логика: AuthService, FileService, ShareService, DocumentSigningService, OpenAIService, JwtService, PasswordService.')
bullet('entities/ -- JPA обекти, директно мапирани към таблици в MySQL.')
bullet('repositories/ -- Spring Data JPA репозитории.')
bullet('utils/ -- SecurityConfig (Spring Security), JobScheduler, GeneralUtils, TokenUtil.')
bullet('config/ -- WebSocketConfig (STOMP/SockJS).')

heading3('2.2.2. Go File Server')

body(
    'Лек микросървис, написан на Go, отговорен единствено за физическото '
    'съхранение на файлове на диска. Приема multipart upload заявки, '
    'запазва файловете под уникално UUID-базирано наименование и ги обслужва за изтегляне. '
    'Spring Boot API делегира качването/изтеглянето на файловете към Go сървъра.'
)

heading3('2.2.3. Next.js Frontend')

body(
    'Потребителският интерфейс е изграден с Next.js 15 (App Router) и React 19. '
    'Ползва TypeScript за типова безопасност и Tailwind CSS v4 за стилизиране. '
    'Всички API заявки се изпращат към Next.js реверс прокси (/proxy/*), '
    'което ги препраща към Spring Boot API за избягване на CORS проблеми.'
)

body('Структура на страниците:')
bullet('/ -- Начална (публична) страница.')
bullet('/login, /register -- Форми за вход и регистрация.')
bullet('/forgot-password, /reset-password -- Управление на парола.')
bullet('/drive -- Основен изглед: всички root папки на потребителя.')
bullet('/drive/folder/[id] -- Съдържание на конкретна папка.')
bullet('/shares -- Управление на всички споделяния.')
bullet('/settings -- Профил и абонамент.')
bullet('/share/[token] -- Публична страница за споделена връзка.')

heading2('2.3. Схема на базата данни')

body(
    'Базата данни се управлява от MySQL 8. Hibernate (JPA) автоматично синхронизира '
    'схемата при стартиране (spring.jpa.hibernate.ddl-auto=update). '
    'По-долу са описани всички основни таблици с техните колони.'
)

heading3('Таблица: users')
add_table(
    ['Колона', 'Тип', 'Описание'],
    [
        ['id', 'BIGINT PK AUTO', 'Уникален идентификатор'],
        ['username', 'VARCHAR UNIQUE NOT NULL', 'Потребителско име (lowercase)'],
        ['email', 'VARCHAR UNIQUE NOT NULL', 'Имейл адрес (lowercase)'],
        ['password', 'VARCHAR NOT NULL', 'BCrypt хеш на паролата'],
        ['createdAt', 'DATETIME NOT NULL', 'Дата на регистрация (автоматична)'],
        ['lastLogin', 'DATETIME', 'Последен успешен вход'],
        ['currentStorageUsed', 'BIGINT NOT NULL', 'Изразходвано пространство в байтове'],
        ['preferredLanguageId', 'BIGINT FK->languages', 'Предпочитан език (опционален)'],
    ]
)

heading3('Таблица: plans')
add_table(
    ['Колона', 'Тип', 'Описание'],
    [
        ['id', 'BIGINT PK AUTO', 'Идентификатор'],
        ['internalName', 'VARCHAR UNIQUE NOT NULL', 'Код на плана: FREE_PLAN, BASIC_PLAN'],
        ['type', 'ENUM(FREE,PAID)', 'Тип на плана'],
        ['storageLimitGB', 'INT NOT NULL', 'Лимит за съхранение в GB'],
        ['pricePerMonth', 'DECIMAL(10,2) NOT NULL', 'Месечна цена'],
        ['currency', 'VARCHAR(3) NOT NULL', 'Валута (напр. BGN)'],
        ['isActive', 'BOOLEAN NOT NULL', 'Дали планът е активен'],
    ]
)

heading3('Таблица: user_subscriptions')
add_table(
    ['Колона', 'Тип', 'Описание'],
    [
        ['id', 'BIGINT PK AUTO', 'Идентификатор'],
        ['user_id', 'BIGINT FK->users UNIQUE', 'Потребител'],
        ['plan_id', 'BIGINT FK->plans', 'Избран план'],
        ['startDate', 'DATETIME NOT NULL', 'Начало на абонамента'],
        ['endDate', 'DATETIME NULL', 'Край (null = активен)'],
        ['status', 'VARCHAR(50) NOT NULL', 'ACTIVE / EXPIRED / CANCELLED'],
        ['renewalDate', 'DATETIME NULL', 'Дата на следващо подновяване'],
        ['paymentMethodId', 'VARCHAR NULL', 'ID на платежен метод в payment gateway'],
    ]
)

heading3('Таблица: folders')
add_table(
    ['Колона', 'Тип', 'Описание'],
    [
        ['id', 'BIGINT PK AUTO', 'Идентификатор'],
        ['canonicalName', 'VARCHAR NOT NULL', 'Базово наименование на папката'],
        ['owner_id', 'BIGINT FK->users NOT NULL', 'Собственик'],
        ['parent_id', 'BIGINT FK->folders NULL', 'Родителска папка (null = root ниво)'],
        ['isDeleted', 'BOOLEAN NOT NULL', 'Soft delete флаг'],
        ['creationDate', 'DATETIME NOT NULL', 'Дата на създаване (автоматична)'],
        ['lastModifiedDate', 'DATETIME NOT NULL', 'Дата на последна промяна (автоматична)'],
    ]
)

heading3('Таблица: files')
add_table(
    ['Колона', 'Тип', 'Описание'],
    [
        ['id', 'BIGINT PK AUTO', 'Идентификатор'],
        ['originalFileName', 'VARCHAR NOT NULL', 'Оригинално наименование на файла'],
        ['uniqueName', 'VARCHAR UNIQUE NOT NULL', 'UUID-базирано уникално физическо наименование'],
        ['type', 'VARCHAR(100) NOT NULL', 'MIME тип (напр. application/pdf)'],
        ['size', 'BIGINT NOT NULL', 'Размер в байтове'],
        ['owner_id', 'BIGINT FK->users NOT NULL', 'Собственик'],
        ['parent_id', 'BIGINT FK->folders NULL', 'Папка (null = root)'],
        ['isDeleted', 'BOOLEAN NOT NULL', 'Soft delete флаг'],
        ['uploadDate', 'DATETIME NOT NULL', 'Дата на качване (автоматична)'],
        ['lastModifiedDate', 'DATETIME NOT NULL', 'Дата на последна промяна'],
    ]
)

heading3('Таблица: shared_items')
add_table(
    ['Колона', 'Тип', 'Описание'],
    [
        ['id', 'BIGINT PK AUTO', 'Идентификатор'],
        ['itemId', 'BIGINT NOT NULL', 'ID на споделения файл/папка'],
        ['itemType', 'ENUM(FILE,FOLDER)', 'Тип на елемента'],
        ['owner_id', 'BIGINT FK->users NOT NULL', 'Потребителят, инициирал споделянето'],
        ['shared_with_user_id', 'BIGINT FK->users NULL', 'Потребител (null = публична връзка)'],
        ['permissionLevel', 'ENUM(VIEW,EDIT,DELETE)', 'Ниво на достъп'],
        ['tokenHash', 'VARCHAR(64) UNIQUE NOT NULL', 'SHA-256 хеш на токена'],
        ['expiresAt', 'DATETIME NULL', 'Дата на изтичане (null = безсрочна)'],
        ['revokedAt', 'DATETIME NULL', 'Дата на оттегляне (null = активно)'],
        ['sharedDate', 'DATETIME NOT NULL', 'Дата на споделяне (автоматична)'],
    ]
)

heading3('Таблица: refresh_tokens')
add_table(
    ['Колона', 'Тип', 'Описание'],
    [
        ['id', 'BIGINT PK AUTO', 'Идентификатор'],
        ['user_id', 'BIGINT FK->users NOT NULL', 'Потребител'],
        ['tokenHash', 'VARCHAR(64) UNIQUE NOT NULL', 'SHA-256 хеш на refresh токена'],
        ['expiresAt', 'TIMESTAMP NOT NULL', 'Дата на изтичане (30 дни)'],
        ['revokedAt', 'TIMESTAMP NULL', 'Дата на оттегляне (null = валиден)'],
        ['createdAt', 'TIMESTAMP NOT NULL', 'Дата на издаване (автоматична)'],
    ]
)

heading3('Таблица: password_reset_tokens')
add_table(
    ['Колона', 'Тип', 'Описание'],
    [
        ['id', 'BIGINT PK AUTO', 'Идентификатор'],
        ['user_id', 'BIGINT FK->users NOT NULL', 'Потребител'],
        ['tokenHash', 'VARCHAR(64) NOT NULL', 'SHA-256 хеш на токена за reset'],
        ['expiresAt', 'DATETIME NOT NULL', 'Изтичане (15 минути след генериране)'],
        ['usedAt', 'DATETIME NULL', 'Дата на употреба (null = неизползван)'],
    ]
)

heading3('Таблица: document_contents')
add_table(
    ['Колона', 'Тип', 'Описание'],
    [
        ['fileId', 'BIGINT PK FK->files', 'ID на файла (1:1 връзка)'],
        ['content', 'LONGTEXT', 'Текстово съдържание на документа'],
        ['lastUpdated', 'DATETIME', 'Момент на последната промяна'],
        ['lastEditorUsername', 'VARCHAR', 'Потребителско им на последния редактор'],
    ]
)

heading3('Таблица: document_signatures')
add_table(
    ['Колона', 'Тип', 'Описание'],
    [
        ['id', 'BIGINT PK AUTO', 'Идентификатор'],
        ['fileId', 'BIGINT FK->files NOT NULL', 'Подписан PDF файл'],
        ['signed_by', 'BIGINT FK->users NOT NULL', 'Потребителят, положил подписа'],
        ['signerName', 'VARCHAR NOT NULL', 'Потребителско им по времето на подписване'],
        ['certificateSubject', 'VARCHAR', 'Subject на X.509 сертификата'],
        ['certificateIssuer', 'VARCHAR', 'Issuer на сертификата'],
        ['certificateValidFrom', 'DATETIME', 'Начало на валидност на сертификата'],
        ['certificateValidTo', 'DATETIME', 'Край на валидност на сертификата'],
        ['signatureLevel', 'VARCHAR', 'Ниво на подписа (PAdES-B-B)'],
        ['signedUniqueName', 'VARCHAR', 'Уникално наименование на подписания PDF файл'],
        ['signedAt', 'DATETIME', 'Момент на подписване'],
    ]
)

diagram_placeholder('database ER diagram')

heading2('2.4. Комуникация между компонентите')

body('Комуникацията между компонентите е организирана по следния начин:')

bullet(
    'Браузър -> Next.js: Потребителят комуникира единствено с Next.js фронтенда '
    'чрез HTTPS (в production среда) или HTTP (development).'
)
bullet(
    'Next.js -> Spring Boot API: Фронтендът изпраща REST заявки към Next.js proxy '
    '(/proxy/*), който ги препраща към http://localhost:8083. Това избягва '
    'CORS проблеми и позволява добавяне на middleware логика.'
)
bullet(
    'Spring Boot API -> MySQL: JPA/Hibernate осъществява заявките към базата данни '
    'чрез JDBC connection pool.'
)
bullet(
    'Spring Boot API -> Go File Server: При качване на файл Spring Boot делегира '
    'физическото запазване на Go сървъра; в текущата имплементация '
    'файловете се записват директно на диска от Spring Boot.'
)
bullet(
    'Браузър <-> Spring Boot API (WebSocket): За реално-времево редактиране '
    'браузърът установява STOMP/SockJS WebSocket връзка директно към /ws.'
)
bullet(
    'Spring Boot API -> OpenAI: ChatController изпраща HTTP заявки към OpenAI REST API '
    'за обработка на чатбот съобщенията (модел gpt-4o-mini).'
)

diagram_placeholder('component communication flow diagram')

heading2('2.5. Сигурност и автентикация')

body(
    'Сигурността е реализирана чрез Spring Security с JWT (JSON Web Token) '
    'автентикация. Архитектурата следва принципа за stateless сесия -- '
    'сървърът не пази сесийно състояние между заявките.'
)

heading3('JWT Authentication Flow')

body('Процесът на автентикация протича в следната последователност:')
bullet('1. Потребителят изпраща POST /api/auth/login с потребителско им/email и парола.')
bullet('2. Spring Boot валидира credentials чрез AuthenticationManager и BCrypt.')
bullet('3. Сървърът издава access token (JWT, валиден 1 час, подписан с HMAC-SHA256) и refresh token (random opaque token, SHA-256 хеширан в БД, валиден 30 дни).')
bullet('4. Фронтендът съхранява двата токена в localStorage.')
bullet('5. При всяка защитена заявка фронтендът изпраща access токена в Authorization: Bearer заглавие.')
bullet('6. При получаване на 401 фронтендът автоматично извиква POST /api/auth/refresh с refresh токена (token rotation -- старият се анулира, нов чифт се издава).')

heading3('Endpoint-и за достъп')
add_table(
    ['Endpoint', 'Метод', 'Достъп', 'Описание'],
    [
        ['/api/auth/**', 'POST', 'Публичен', 'Регистрация, вход, refresh, logout, reset парола'],
        ['/public/shares/**', 'GET', 'Публичен', 'Публични споделени връзки'],
        ['/ws/**', 'WS', 'Публичен', 'WebSocket endpoint (STOMP/SockJS)'],
        ['/api/**', 'ALL', 'Authenticated (JWT)', 'Всички останали API операции'],
    ]
)

heading2('2.6. Инфраструктура -- Docker и Nginx')

heading3('2.6.1. Docker и Docker Compose')

body(
    'MyDrive е проектирана за разгръщане с Docker и Docker Compose. '
    'Контейнеризацията осигурява изолация на компонентите, '
    'повторяемост на средата (development, testing, production) '
    'и лесна мащабируемост.'
)

body('Основни предимства на Docker в контекста на MyDrive:')
bullet('Изолация -- всеки сервиз работи в собствен контейнер с точно дефинирани зависимости.')
bullet('Повторяемост -- едно и също Docker изображение дава идентично поведение навсякъде.')
bullet('Зависимости -- MySQL се стартира преди Spring Boot чрез depends_on с healthcheck.')
bullet('Волуми -- файловото хранилище и базата данни се съхраняват в именувани Docker волуми, независими от контейнера.')

code_block('''version: "3.9"
services:
  mysql:
    image: mysql:8.0
    environment:
      MYSQL_DATABASE: mydrivedb
      MYSQL_ROOT_PASSWORD: root
    volumes:
      - mysql_data:/var/lib/mysql
    healthcheck:
      test: ["CMD", "mysqladmin", "ping", "-h", "localhost"]
      interval: 10s
      retries: 5

  backend:
    build: ./02_Development/backend/MYDrive
    ports:
      - "8083:8083"
    depends_on:
      mysql:
        condition: service_healthy
    environment:
      SPRING_DATASOURCE_URL: jdbc:mysql://mysql:3306/mydrivedb
      APP_JWT_SECRET: ${JWT_SECRET}
      APP_STORAGE_PATH: /storage
    volumes:
      - file_storage:/storage

  file-server:
    build: ./02_Development/backend/file-server
    ports:
      - "8080:8080"
    volumes:
      - file_storage:/storage

  frontend:
    build: ./02_Development/frontend/web
    ports:
      - "3000:3000"
    environment:
      NEXT_PUBLIC_API_URL: /proxy

  nginx:
    image: nginx:alpine
    ports:
      - "80:80"
      - "443:443"
    volumes:
      - ./nginx/nginx.conf:/etc/nginx/nginx.conf:ro
      - ./nginx/ssl:/etc/nginx/ssl:ro
    depends_on:
      - backend
      - frontend

volumes:
  mysql_data:
  file_storage:''')

heading3('2.6.2. Nginx -- Реверс прокси и SSL терминация')

body(
    'Nginx служи като единствена точка за входящ трафик (single entry point). '
    'Основните функции, които изпълнява в системата MyDrive, са:'
)
bullet('SSL/TLS терминация -- Nginx декриптира HTTPS трафика и го препраща до вътрешните сервизи по HTTP.')
bullet('Реверс прокси -- маршрутизира /api/* и /public/* заявки към Spring Boot (порт 8083), а останалото -- към Next.js (порт 3000).')
bullet('WebSocket проксиране -- специална конфигурация за Upgrade и Connection заглавия за STOMP/SockJS.')
bullet('Сигурност -- добавя HTTP security headers (HSTS, X-Frame-Options, X-Content-Type-Options) и ограничава директен достъп до вътрешните портове.')
bullet('Пренасочване HTTP -> HTTPS -- всички незашифровани заявки се пренасочват към HTTPS.')

code_block('''# nginx/nginx.conf (примерна конфигурация)
events { worker_connections 1024; }

http {
  upstream backend  { server backend:8083; }
  upstream frontend { server frontend:3000; }

  # HTTP -> HTTPS redirect
  server {
    listen 80;
    server_name mydrive.example.com;
    return 301 https://$host$request_uri;
  }

  server {
    listen 443 ssl;
    server_name mydrive.example.com;

    ssl_certificate     /etc/nginx/ssl/fullchain.pem;
    ssl_certificate_key /etc/nginx/ssl/privkey.pem;
    ssl_protocols       TLSv1.2 TLSv1.3;
    ssl_ciphers         HIGH:!aNULL:!MD5;

    # Security headers
    add_header Strict-Transport-Security "max-age=31536000" always;
    add_header X-Frame-Options DENY;
    add_header X-Content-Type-Options nosniff;

    # API и публични endpoints -> Spring Boot
    location ~ ^/(api|public)/ {
      proxy_pass http://backend;
      proxy_set_header Host $host;
      proxy_set_header X-Real-IP $remote_addr;
      proxy_set_header X-Forwarded-For $proxy_add_x_forwarded_for;
    }

    # WebSocket (STOMP/SockJS) -> Spring Boot
    location /ws {
      proxy_pass http://backend;
      proxy_http_version 1.1;
      proxy_set_header Upgrade $http_upgrade;
      proxy_set_header Connection "upgrade";
      proxy_set_header Host $host;
    }

    # Всичко останало -> Next.js Frontend
    location / {
      proxy_pass http://frontend;
      proxy_set_header Host $host;
      proxy_set_header X-Real-IP $remote_addr;
    }
  }
}''')

diagram_placeholder('Docker and Nginx deployment architecture')

page_break()

# ==============================================================================
# CHAPTER 3
# ==============================================================================

heading1('3. ОПИСАНИЕ НА КОДА')

heading2('3.1. Backend -- Spring Boot')

heading3('3.1.1. Стартиране на приложението')

body(
    'Входната точка на бекенда е класът MyDriveApplication.java, анотиран с '
    '@SpringBootApplication. При стартиране Spring Boot инициализира всички Spring beans, '
    'стартира вградения Tomcat сървър на порт 8083 и Hibernate синхронизира '
    'схемата на базата данни.'
)

body(
    'Конфигурационните параметри се намират в src/main/resources/application.properties: '
    'connection string към MySQL, JWT secret ключ и expiration time, '
    'директория за физическо съхранение на файлове, SMTP настройки за имейли '
    'и пътища към PKCS#12 keystore за eIDAS подписване.'
)

heading3('3.1.2. Автентикация -- AuthController и AuthService')

body(
    'AuthController (пакет controllers/) дефинира REST endpoint-и за целия '
    'жизнен цикъл на автентикацията. '
    'AuthService съдържа бизнес логиката за всяка операция.'
)

body('Основни методи на AuthService:')
bullet('register() -- валидира уникалност на email/username, хешира паролата с BCrypt, създава потребителя и го записва автоматично в FREE абонаментен план.')
bullet('login() -- автентикира чрез Spring Security AuthenticationManager, издава JWT access token (1 ч.) и opaque refresh token (30 дни, SHA-256 хеширан в БД).')
bullet('refresh() -- проверява refresh токена в БД, извършва token rotation.')
bullet('logout() -- анулира refresh токена в базата.')
bullet('forgotPassword() / resetPassword() -- изпраща имейл с токен, валиден 15 минути; reset-ва паролата след валидация.')

code_block('''// AuthController.java -- фрагмент
@PostMapping("/login")
public AuthTokensResponse login(@Valid @RequestBody UserLoginRequest req) {
    return authService.login(req);
}

@PostMapping("/refresh")
public AuthTokensResponse refresh(@Valid @RequestBody RefreshRequest req) {
    return authService.refresh(req);  // token rotation
}

@GetMapping("/me")
public UserResponse me(JwtAuthenticationToken auth) {
    Long userId = ((Number) auth.getTokenAttributes().get("uid")).longValue();
    return authService.getUserResponseById(userId);
}''')

body(
    'JwtService генерира JWT токена с claim-ове sub (email) и uid (userId), '
    'подписан с HMAC-SHA256 ключ, конфигуриран в application.properties (app.jwt.secret). '
    'Токенът се валидира автоматично от Spring Security OAuth2 Resource Server '
    'при всяка защитена заявка.'
)

heading3('3.1.3. Управление на файлове -- FileController и FileService')

body(
    'FileController дефинира CRUD endpoint-и за файлове (/api/files). '
    'FileService управлява физическото запазване и метаданните в БД.'
)

body('Основни операции:')
bullet('upload() -- приема MultipartFile, генерира UUID-базирано уникално наименование, записва файла на диска и създава запис в таблица files.')
bullet('getDownloadResource() -- връща PathResource за изтегляне с оригиналното наименование и MIME тип.')
bullet('update() -- преименува файла (само метаданни в БД, физическият файл не се мести).')
bullet('move() -- мести файл в друга папка (промяна на parent FK в БД).')
bullet('delete() -- soft delete (isDeleted=true) + физическо изтриване от диска.')

code_block('''// FileService.java -- upload метод (фрагмент)
public FileResponse upload(MultipartFile file, Long parentFolderId, Long ownerId)
        throws IOException {
    String uniqueName = UUID.randomUUID() + "_" + file.getOriginalFilename();
    Path target = Path.of(storagePath, uniqueName);
    Files.copy(file.getInputStream(), target);  // physical write

    FileEntity entity = FileEntity.builder()
        .originalFileName(file.getOriginalFilename())
        .uniqueName(uniqueName)
        .type(file.getContentType())
        .size(file.getSize())
        .owner(owner)
        .parent(parent)
        .isDeleted(false)
        .build();

    return FileResponse.from(fileRepository.save(entity));
}''')

diagram_placeholder('file upload sequence diagram')

heading3('3.1.4. Управление на папки -- FolderController и FolderService')

body(
    'Папките поддържат самореферентна йерархия (parent_id FK към същата таблица folders). '
    'Root папките имат parent=null. '
    'Папките поддържат каскадно изтриване -- при soft delete на папка '
    'всички нейни деца се маркират като изтрити рекурсивно.'
)

body('Основни операции:')
bullet('list(parentId) -- листване на всички папки за текущото ниво (null = root).')
bullet('create(canonicalName, parentId) -- създаване на нова папка.')
bullet('update(id, canonicalName) -- преименуване.')
bullet('delete(id) -- soft delete с каскадно разпространение към дъщерните елементи.')

heading3('3.1.5. Споделяне -- ShareController и ShareService')

body(
    'Споделянето работи чрез генерирани токен-връзки. '
    'Механизмът е проектиран с оглед на сигурността -- '
    'raw токенът никога не се записва в базата данни.'
)

body('Процес на споделяне:')
bullet('1. Генерира се случаен opaque token (32 байта, криптографски сигурен генератор).')
bullet('2. Токенът се SHA-256 хешира -- само хешът се пази в колона tokenHash.')
bullet('3. Генерираният raw токен се връща веднъж на потребителя.')
bullet('4. При публичен достъп (/public/shares/{token}) raw токенът се хешира и се търси в БД.')
bullet('5. Проверяват се: дали е оттеглен (revokedAt != null) и дали е изтекъл (expiresAt).')

code_block('''// ShareService.java -- create метод (фрагмент)
String rawToken  = TokenUtil.newOpaqueToken();     // 32 random bytes, URL-safe Base64
String tokenHash = TokenUtil.sha256Hex(rawToken);  // stored in DB

SharedItemEntity share = SharedItemEntity.builder()
    .owner(owner)
    .itemId(req.itemId())
    .itemType(itemType)
    .permissionLevel(perm)
    .tokenHash(tokenHash)   // only hash is persisted
    .expiresAt(req.expiresAt())
    .build();

sharedItemRepository.save(share);
// raw token returned once -- never stored
return new ShareCreateResponse(saved.getId(), rawToken, ...);''')

diagram_placeholder('sharing token flow diagram')

heading3('3.1.6. Spring Security конфигурация')

body(
    'SecurityConfig конфигурира Spring Security за stateless режим: '
    'CSRF е изключен (REST API не използва cookie-базирани сесии), '
    'а JWT валидацията се осъществява автоматично от Spring Security OAuth2 Resource Server.'
)

code_block('''// SecurityConfig.java
http.csrf(AbstractHttpConfigurer::disable)
    .sessionManagement(sm -> sm
        .sessionCreationPolicy(SessionCreationPolicy.STATELESS))
    .authorizeHttpRequests(auth -> auth
        .requestMatchers("/api/auth/**").permitAll()
        .requestMatchers("/public/shares/**").permitAll()
        .requestMatchers("/ws/**").permitAll()
        .anyRequest().authenticated())
    .oauth2ResourceServer(oauth -> oauth.jwt(Customizer.withDefaults()));''')

heading2('3.2. File Server -- Go')

body(
    'Go микросървисът (main.go) е минималистичен HTTP сървър, '
    'отговарящ единствено за физическото боравене с файлове. '
    'Реализирани са два основни endpoint-а:'
)
bullet('POST /upload -- получава multipart форма, записва файла под UUID-базирано наименование в настроена директория.')
bullet('GET /download/{filename} -- обслужва файла по уникалното му наименование с правилни HTTP заглавия.')

body(
    'Разделянето на файловия сървър в отделен микросървис осигурява '
    'независимо мащабиране на хранилището спрямо бизнес логиката. '
    'В production среда Go File Server може да се замени с S3-съвместимо хранилище '
    '(MinIO, AWS S3, Cloudflare R2) без промяна в Spring Boot логиката.'
)

heading2('3.3. Frontend -- Next.js')

heading3('3.3.1. Структура и маршрутизиране')

body(
    'Фронтендът следва App Router модела на Next.js 15. '
    'Защитените страници са в директория (protected)/, '
    'чийто layout.tsx проверява наличието на access token -- '
    'при липса пренасочва към /login. '
    'Публичните страници (/, /login, /register, /share/[token]) '
    'са достъпни без автентикация.'
)

heading3('3.3.2. AuthContext -- глобално управление на автентикацията')

body(
    'AuthContext (contexts/AuthContext.tsx) управлява глобалното автентикационно '
    'състояние на приложението чрез React Context API. '
    'При зареждане прочита кешираните user и accessToken от localStorage '
    'и валидира сесията чрез GET /api/auth/me.'
)

code_block('''// AuthContext.tsx -- инициализация при зареждане
useEffect(() => {
  const cached = localStorage.getItem("user");
  const token  = localStorage.getItem("accessToken");
  if (cached && token) {
    setUser(JSON.parse(cached));
    refreshUser().finally(() => setLoading(false));
  } else {
    setLoading(false);
  }
}, [refreshUser]);''')

heading3('3.3.3. API клиент (lib/api.ts)')

body(
    'Централният API клиент (lib/api.ts) капсулира всички HTTP заявки '
    'в типизирани функции. Функцията request<T>() автоматично:'
)
bullet('Добавя Authorization: Bearer {accessToken} заглавие при всяка заявка.')
bullet('При получаване на 401 извиква refreshTokens() и повтаря заявката веднъж.')
bullet('При неуспешен refresh изчиства токените и пренасочва към /login.')

body('Дефинирани API групи:')
bullet('auth -- login, register, logout, me, forgotPassword, resetPassword, changePassword.')
bullet('filesApi -- list, get, create, update, move, delete, upload (multipart), download.')
bullet('foldersApi -- list, create, update, delete.')
bullet('sharesApi -- list, create, revoke.')
bullet('publicSharesApi -- get (без автентикация).')
bullet('chatApi -- send (изпраща plain text, получава plain text отговор).')
bullet('documentsApi -- getContent (последно съдържание на документ).')
bullet('signingApi -- sign (eIDAS PAdES), getSignatures, downloadUrl.')

heading3('3.3.4. DriveView -- централен компонент на файловия мениджър')

body(
    'DriveView (components/drive/DriveView.tsx) е главният компонент '
    'за управление на файловата структура. Зарежда папките и файловете '
    'за текущото ниво, показва ги в грид изглед и управлява всички модали:'
)

bullet('CreateFolderModal -- создаване на нова папка.')
bullet('UploadModal -- качване на файл с multipart форма.')
bullet('RenameModal -- преименуване на файл или папка.')
bullet('DeleteModal -- потвърждение за изтриване.')
bullet('ShareModal -- генериране на споделена връзка с настройки за достъп.')
bullet('MoveModal -- местене на файл в друга папка.')
bullet('DocumentEditor -- вграден реално-времеви текстов редактор.')
bullet('SignDocumentModal -- eIDAS PAdES подписване на PDF.')

diagram_placeholder('DriveView component tree')

heading3('3.3.5. Next.js Proxy Rewrite')

body(
    'В next.config.ts е конфигуриран rewrite: '
    'всички заявки към /proxy/* се препращат към http://localhost:8083/:path*. '
    'Това позволява фронтендът да изпраща API заявки относително (/proxy/api/...) '
    'без нужда от CORS конфигурация на бекенда.'
)

code_block('''// next.config.ts
const nextConfig = {
  async rewrites() {
    return [{
      source: "/proxy/:path*",
      destination: "http://localhost:8083/:path*",
    }];
  },
};''')

heading2('3.4. Реално-времево съвместно редактиране (WebSocket / STOMP)')

body(
    'MyDrive поддържа едновременно редактиране на текстови документи '
    'от множество потребители чрез WebSocket (STOMP протокол над SockJS fallback).'
)

heading3('3.4.1. Сървърна конфигурация (WebSocketConfig)')

body(
    'WebSocketConfig конфигурира STOMP message broker: '
    '/topic/* за broadcast канали (публикуване към всички абонирани клиенти), '
    '/app/* за application destinations (съобщения, обработвани от контролери). '
    'SockJS е активиран като fallback за браузъри без нативна WebSocket поддръжка.'
)

code_block('''// WebSocketConfig.java
@Override
public void configureMessageBroker(MessageBrokerRegistry config) {
    config.enableSimpleBroker("/topic");        // broadcast destinations
    config.setApplicationDestinationPrefixes("/app");  // controller destinations
}

@Override
public void registerStompEndpoints(StompEndpointRegistry registry) {
    registry.addEndpoint("/ws")
            .setAllowedOriginPatterns("*")
            .withSockJS();  // SockJS fallback
}''')

heading3('3.4.2. CollaborationController')

body(
    'CollaborationController получава редакторски съобщения на '
    '@MessageMapping("/document/{fileId}/edit"), '
    'записва последното съдържание в таблица document_contents '
    'и го broadcast-ва към @SendTo("/topic/document/{fileId}") -- '
    'всички абонирани за тази тема клиенти получават съобщението незабавно.'
)

code_block('''// CollaborationController.java
@MessageMapping("/document/{fileId}/edit")
@SendTo("/topic/document/{fileId}")
public DocumentEditMessage broadcastEdit(
        @DestinationVariable Long fileId,
        DocumentEditMessage message) {

    // Persist latest snapshot in document_contents table
    DocumentContentEntity entity = contentRepository.findById(fileId)
        .orElseGet(() -> { /* create new */ });
    entity.setContent(message.content());
    entity.setLastEditorUsername(message.editorUsername());
    contentRepository.save(entity);

    return message;  // broadcast to all subscribers on /topic/document/{fileId}
}''')

heading3('3.4.3. Клиентска страна (DocumentEditor)')

body(
    'DocumentEditor (components/drive/DocumentEditor.tsx) '
    'установява STOMP/SockJS връзка, абонира се за /topic/document/{fileId} '
    'и изпраща съобщения при всяка промяна в текстовото поле чрез /app/document/{fileId}/edit. '
    'Показва потребителското им на последния редактор в реално време.'
)

diagram_placeholder('WebSocket collaboration sequence diagram')

heading2('3.5. Електронно подписване на документи (eIDAS PAdES)')

body(
    'MyDrive реализира PAdES-B-B (PDF Advanced Electronic Signature -- Baseline-B) '
    'подписване, съвместимо с Регламент eIDAS (EU 910/2014), Член 26 -- '
    'Advanced Electronic Signature (AdES).'
)

heading3('3.5.1. Автоматично генериране на demo keystore')

body(
    'При първо стартиране DocumentSigningService.initKeystore() '
    '(анотиран с @PostConstruct) проверява дали съществува keystore файл (.p12). '
    'Ако не -- автоматично генерира:'
)
bullet('RSA-2048 key pair.')
bullet('Самоподписан X.509 v3 сертификат (Subject: CN=MyDrive Demo Signature, O=MyDrive, C=BG, валиден 10 години).')
bullet('PKCS#12 (.p12) файл, съдържащ private key и сертификата.')

body(
    'Пътят до keystore файла се конфигурира в application.properties '
    '(app.signing.keystore-path, app.signing.keystore-password, app.signing.key-alias).'
)

heading3('3.5.2. Процес на PAdES-B-B подписване')

body('Подписването на PDF файл протича в следните последователни стъпки:')
bullet('1. Зареждане на оригиналния PDF с Apache PDFBox 3.x (Loader.loadPDF()).')
bullet('2. Създаване на PDSignature обект с FILTER_ADOBE_PPKLITE и SUBFILTER_ADBE_PKCS7_DETACHED.')
bullet('3. Генериране на CMS (Cryptographic Message Syntax) подпис с BouncyCastle -- SHA256withRSA.')
bullet('4. Вграждане на PKCS#7 подписа в PDF чрез incrementalSave (incremental update -- оригиналното съдържание остава непроменено).')
bullet('5. Записване на подписания файл на диска под ново уникално наименование (signed_{uuid}.pdf).')
bullet('6. Съхраняване на метаданни за подписа в таблица document_signatures.')

code_block('''// DocumentSigningService.java -- signPdf метод (фрагмент)
try (PDDocument doc = Loader.loadPDF(sourceFile.toFile())) {
    PDSignature signature = new PDSignature();
    signature.setFilter(PDSignature.FILTER_ADOBE_PPKLITE);
    signature.setSubFilter(PDSignature.SUBFILTER_ADBE_PKCS7_DETACHED);
    signature.setName(user.getUsername());
    signature.setReason("PAdES-B-B eIDAS Advanced Electronic Signature");

    doc.addSignature(signature, (InputStream content) -> {
        byte[] bytes = content.readAllBytes();
        CMSSignedDataGenerator gen = new CMSSignedDataGenerator();
        ContentSigner cs = new JcaContentSignerBuilder("SHA256withRSA")
            .setProvider("BC").build(privateKey);
        gen.addSignerInfoGenerator(...);
        gen.addCertificates(new JcaCertStore(List.of(chain)));
        CMSSignedData cmsData = gen.generate(
            new CMSProcessableByteArray(bytes), false);
        return cmsData.getEncoded();
    }, opts);

    doc.saveIncremental(Files.newOutputStream(signedFile));
}''')

body(
    'В производствена среда demo сертификатът се заменя с такъв от '
    'квалифициран доставчик на удостоверителни услуги (QSP) по смисъла на eIDAS, '
    'за постигане на QES (Qualified Electronic Signature) ниво.'
)

diagram_placeholder('eIDAS PDF signing flow')

heading2('3.6. AI Чатбот')

body(
    'MyDrive интегрира AI асистент, реализиран като плаваща widget в долния десен ъгъл. '
    'ChatBot.tsx компонентът изпраща съобщенията към ChatController '
    '(POST /api/chat, Content-Type: text/plain). '
    'Сървърният ChatController делегира към OpenAIService, '
    'който извиква OpenAI Chat Completions API (gpt-4o-mini модел) '
    'и връща отговора като plain text.'
)

body(
    'Чатботът може да отговаря на въпроси, свързани с ползването на платформата, '
    'и е достъпен от всички защитени страници.'
)

diagram_placeholder('AI chatbot widget and flow')

page_break()

# ==============================================================================
# CHAPTER 4
# ==============================================================================

heading1('4. РЪКОВОДСТВО ЗА ПОТРЕБИТЕЛЯ')

body(
    'Настоящото ръководство описва стъпка по стъпка основните операции в системата MyDrive. '
    'За достъп е необходим модерен браузър (Chrome, Firefox, Edge, Safari) '
    'и активна интернет връзка.'
)

heading2('4.1. Регистрация и вход в системата')

heading3('4.1.1. Регистрация на нов акаунт')

body('За да създадете нов акаунт в MyDrive:')
bullet('Отворете началната страница на MyDrive в браузъра.')
bullet('Натиснете бутона Регистрация или отидете директно на /register.')
bullet('Попълнете полетата: Потребителско им, Имейл адрес и Парола.')
bullet('Натиснете "Register".')
bullet('При успех ще бъдете автоматично вписани и пренасочени към файловия мениджър (/drive).')

body(
    'Забележка: Потребителското им и имейлът трябва да са уникални в системата. '
    'При регистрация автоматично се активира безплатен (FREE) план с лимитирано пространство.'
)

diagram_placeholder('registration page screenshot')

heading3('4.1.2. Вход в системата')

body('За вход в съществуващ акаунт:')
bullet('Отидете на /login.')
bullet('Въведете потребителско им или имейл и паролата си.')
bullet('Натиснете "Login".')
bullet('Системата ще ви пренасочи към файловия мениджър (/drive).')

body(
    'Ако сте забравили паролата си, натиснете "Forgot password" и '
    'следвайте инструкциите. Ще получите имейл с връзка за нулиране, '
    'валидна 15 минути.'
)

diagram_placeholder('login page screenshot')

heading2('4.2. Управление на папки и файлове')

heading3('4.2.1. Навигация')

body(
    'След вход попадате в главния изглед на My Drive (/drive), '
    'показващ всички ваши root папки. '
    'Натиснете върху папка за навигация в нея. '
    'В горната part на страницата се показва breadcrumb навигация '
    'за лесно придвижване назад.'
)

heading3('4.2.2. Създаване на нова папка')

body('За да създадете нова папка:')
bullet('Натиснете бутона "New folder" (иконата с папка+) в лентата с инструменти.')
bullet('Въведете им на папката в диалоговия прозорец.')
bullet('Натиснете "Create".')
bullet('Новата папка ще се появи незабавно в текущия изглед.')

diagram_placeholder('create folder modal screenshot')

heading3('4.2.3. Преименуване')

body(
    'Задръжте курсора над папка или файл и натиснете иконата за опции (три точки). '
    'Изберете "Rename", въведете новото наименование и потвърдете.'
)

heading3('4.2.4. Изтриване')

body(
    'Изберете "Delete" от контекстното меню. '
    'Ще бъдете помолени за потвърждение. '
    'Изтритите файлове се маркират като изтрити (soft delete) '
    'и физически се премахват от диска.'
)

heading3('4.2.5. Местене на файл')

body(
    'От контекстното меню на файл изберете "Move". '
    'В диалоговия прозорец изберете целевата папка от списъка. '
    'Натиснете "Move" за потвърждение.'
)

heading2('4.3. Качване и изтегляне на файлове')

heading3('4.3.1. Качване на файл')

body(
    'Файловете могат да се качват само вътре в папка (не на root ниво). '
    'Влезте в желаната папка и:'
)
bullet('Натиснете бутона "Upload" в лентата с инструменти.')
bullet('В диалоговия прозорец изберете файла от вашия компютър.')
bullet('Натиснете "Upload" за потвърждение.')
bullet('При успех файлът ще се появи в текущата папка.')

body(
    'Поддържаните формати са всички MIME типове. '
    'Максималният капацитет зависи от абонаментния план.'
)

diagram_placeholder('file upload modal screenshot')

heading3('4.3.2. Изтегляне на файл')

body(
    'Натиснете иконата за опции (три точки) на желания файл '
    'и изберете "Download". '
    'Файлът ще бъде изтеглен с оригиналното си наименование.'
)

heading2('4.4. Споделяне на файлове и папки')

body(
    'MyDrive позволява споделяне на файлове и папки чрез генерирани токен-връзки. '
    'Споделянето може да е с определен краен срок и ниво на достъп (VIEW или EDIT).'
)

heading3('4.4.1. Създаване на споделена връзка')

bullet('Натиснете "Share" от контекстното меню на файл или папка.')
bullet('В модалния прозорец изберете: Ниво на достъп (VIEW или EDIT) и (опционално) дата на изтичане.')
bullet('Натиснете "Create share link".')
bullet('Копирайте генерирания URL и го изпратете на получателя.')

body(
    'Важно: Raw токенът се показва само веднъж. '
    'В базата данни се пази само SHA-256 хешът. '
    'Ако изгубите връзката, трябва да генерирате нова.'
)

diagram_placeholder('share modal screenshot')

heading3('4.4.2. Управление на споделяния')

body(
    'На страница /shares можете да видите всички ваши активни и оттеглени споделяния. '
    'За всяко споделяне се показват: тип (FILE/FOLDER), ниво на достъп, '
    'дата на създаване, дата на изтичане и статус.'
)

body(
    'За да оттеглите споделяне, натиснете "Revoke" до желаното споделяне. '
    'Оттеглените връзки спират да работят незабавно.'
)

heading3('4.4.3. Достъп до споделена връзка')

body(
    'Получателят на споделена връзка може да отвори файла/папката на /share/{token} '
    'без нужда от акаунт в системата (публичен достъп). '
    'Показват се: тип на елемента, ниво на достъп и дата на изтичане.'
)

heading2('4.5. Съвместно редактиране на документи')

body(
    'MyDrive поддържа реално-времево съвместно редактиране на текстови документи '
    'от множество потребители едновременно чрез WebSocket технология.'
)

heading3('4.5.1. Отваряне на документ за редактиране')

bullet('Намерете текстов файл в папката.')
bullet('Натиснете "Edit" от контекстното меню.')
bullet('Ще се отвори вграденият редактор (DocumentEditor).')

heading3('4.5.2. Работа в редактора')

body(
    'Редакторът показва текущото съдържание на документа. '
    'Всяка промяна, която направите, незабавно се синхронизира с всички '
    'потребители, отворили същия документ в момента. '
    'В горната part на редактора се вижда потребителското им на последния редактор.'
)

body(
    'При затваряне на редактора съдържанието се запазва автоматично. '
    'При следващо отваряне се зарежда последното записано съдържание.'
)

diagram_placeholder('collaborative document editor screenshot')

heading2('4.6. Електронно подписване на PDF')

body(
    'MyDrive поддържа PAdES-B-B електронен подпис за PDF документи, '
    'съвместим с европейския регламент eIDAS.'
)

heading3('4.6.1. Подписване на PDF документ')

bullet('Намерете PDF файл в папката.')
bullet('Натиснете "Sign" от контекстното меню.')
bullet('В модалния прозорец (SignDocumentModal) ще видите информация за файла.')
bullet('Натиснете "Sign document" за потвърждение.')
bullet('Системата ще генерира подписан PDF файл и ще покаже детайлите за подписа.')

heading3('4.6.2. Преглед и изтегляне на подписания документ')

body(
    'След успешно подписване в модала ще видите информация за подписа: '
    'подписващ (signerName), сертификат Subject/Issuer, дата на подписване '
    'и ниво на подписа (PAdES-B-B). '
    'Можете да изтеглите подписания PDF чрез бутона "Download signed PDF".'
)

body(
    'Забележка: Само PDF файлове могат да се подписват. '
    'За другите формати опцията "Sign" не е налична.'
)

diagram_placeholder('sign document modal screenshot')

heading2('4.7. AI Асистент')

body(
    'На всяка защитена страница в долния десен ъгъл се показва иконата на AI асистента. '
    'Натиснете я, за да отворите чат прозореца.'
)

body('Можете да задавате въпроси на естествен език, например:')
bullet('"Как да споделя папка с колега?"')
bullet('"Как работи електронното подписване?"')
bullet('"Какво е PAdES?"')

body(
    'AI асистентът е базиран на GPT-4o mini и може да отговаря '
    'на широк спектър от въпроси, свързани с ползването на платформата.'
)

diagram_placeholder('AI chatbot widget screenshot')

heading2('4.8. Управление на профила и абонамента')

body('Отидете на /settings за достъп до настройките на профила.')

heading3('4.8.1. Промяна на парола')

bullet('В секция Security въведете текущата парола и новата парола.')
bullet('Натиснете "Change password".')
bullet('Паролата ще бъде сменена незабавно.')

heading3('4.8.2. Информация за абонамента')

body(
    'В секция Subscription се показва текущият абонаментен план, '
    'лимитът за съхранение и изразходваното пространство '
    'чрез визуална лента (StorageBar). '
    'За промяна на плана се свържете с администратора на системата.'
)

diagram_placeholder('settings page and storage bar screenshot')

page_break()

# ==============================================================================
# CONCLUSION
# ==============================================================================

heading1('ЗАКЛЮЧЕНИЕ')

body(
    'Настоящата дипломна работа представя проектирането и реализацията на MyDrive -- '
    'пълнофункционална, self-hosted облачна платформа за управление на документи. '
    'Системата обединява съвременни технологии в кохерентна архитектура: '
    'Spring Boot 3 с JWT автентикация и eIDAS PAdES подписване, '
    'Next.js 15 с React 19 за потребителския интерфейс, '
    'MySQL за релационно съхранение на метаданни, '
    'STOMP/WebSocket за реално-времево сътрудничество '
    'и AI асистент, базиран на GPT-4o mini.'
)

body(
    'Инфраструктурата е проектирана с Docker Compose и Nginx, '
    'осигурявайки изолация на компонентите, повторяемост на средата '
    'и SSL терминация. '
    'Архитектурата с разделени отговорности позволява независимо мащабиране '
    'на бекенда, файловия сървър и фронтенда.'
)

body(
    'Механизмът за сигурност следва добрите практики: '
    'BCrypt за хеширане на пароли, stateless JWT автентикация с token rotation, '
    'SHA-256 хеширане на refresh и share токени преди запис в БД, '
    'Spring Security с OAuth2 Resource Server за валидация на JWT.'
)

body(
    'В бъдещи версии на системата се планира: '
    'мобилно приложение (React Native), '
    'интеграция с квалифициран TSP за QES (Qualified Electronic Signature) подписване, '
    'платежна интеграция за абонаментните планове, '
    'CRDT-базирано разрешаване на конфликти при едновременно редактиране '
    'и versioning на документи.'
)

# ==============================================================================
# SAVE
# ==============================================================================

out = '/home/slavi/uni/diplomna/MyDrive/MyDrive_Dokumentaciq.docx'
doc.save(out)
print('Saved: ' + out)
